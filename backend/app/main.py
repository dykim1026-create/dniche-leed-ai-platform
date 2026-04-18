from datetime import datetime
from pathlib import Path
from uuid import uuid4
import json
import csv
from io import StringIO

from fastapi import FastAPI, Depends, HTTPException, UploadFile, File, Query
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from sqlalchemy import text
from sqlalchemy.orm import Session
from pypdf import PdfReader
from docx import Document as DocxDocument
from openpyxl import load_workbook

from app.db import Base, engine, SessionLocal
from app.models import Project, Document
from app.schemas import (
    ProjectCreate,
    ProjectResponse,
    DocumentResponse,
    DocumentSearchResultResponse,
    ReviewEvidenceItemResponse,
    CorrectiveActionResponse,
    ReviewFindingResponse,
    Agent1ReviewResponse,
    EnergyFindingResponse,
    Agent2EnergyReviewResponse,
    CarbonFindingResponse,
    Agent3CarbonReviewResponse,
    LeedScoringFindingResponse,
    Agent4LeedScoringResponse,
)

app = FastAPI(title="Dniche LEED AI Backend")

UPLOAD_DIR = Path("/app/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

TEXT_EXTENSIONS = {".txt", ".md", ".csv", ".json"}
SUPPORTED_PARSE_EXTENSIONS = {".txt", ".md", ".csv", ".json", ".pdf", ".docx", ".xlsx"}

REVIEW_TOPICS = [
    {
        "topic_id": "energy_performance",
        "topic_name": "Energy Performance",
        "keywords": ["energy", "eui", "hvac", "lighting", "envelope", "ashrae"],
        "recommendation": "Add or improve energy model evidence, HVAC narrative, lighting power details, and envelope performance documentation.",
    },
    {
        "topic_id": "water_efficiency",
        "topic_name": "Water Efficiency",
        "keywords": ["water", "fixture", "flow rate", "irrigation", "gpm", "gpf"],
        "recommendation": "Add fixture schedules, water use calculations, irrigation strategy, and indoor/outdoor water reduction evidence.",
    },
    {
        "topic_id": "materials",
        "topic_name": "Materials and Embodied Carbon",
        "keywords": ["material", "epd", "recycled", "concrete", "steel", "carbon"],
        "recommendation": "Add material schedules, EPD references, recycled content information, and carbon-related documentation.",
    },
    {
        "topic_id": "ieq",
        "topic_name": "Indoor Environmental Quality",
        "keywords": ["ventilation", "voc", "daylight", "co2", "thermal comfort", "iaq"],
        "recommendation": "Add ventilation basis, IAQ strategy, VOC-related specifications, daylight evidence, and thermal comfort notes.",
    },
]

CORRECTIVE_ACTION_LIBRARY = {
    "energy_performance": [
        {
            "discipline": "Architecture",
            "action": "Provide envelope performance summary including wall, roof, glazing, and shading assumptions.",
            "reason": "Envelope-related evidence is needed to support energy compliance review.",
        },
        {
            "discipline": "Mechanical",
            "action": "Provide HVAC system description, efficiencies, controls sequence, and ventilation basis.",
            "reason": "Mechanical design evidence is central to energy performance assessment.",
        },
        {
            "discipline": "Electrical",
            "action": "Provide lighting power density, lighting controls, and major equipment load assumptions.",
            "reason": "Lighting and connected loads affect energy evaluation and documentation.",
        },
        {
            "discipline": "Sustainability",
            "action": "Prepare or update the energy model narrative and cross-check design inputs against LEED submission needs.",
            "reason": "A coordinated sustainability review is needed to consolidate evidence.",
        },
    ],
    "water_efficiency": [
        {
            "discipline": "Plumbing",
            "action": "Provide plumbing fixture schedule with flow rates and flush rates for all fixture types.",
            "reason": "Fixture performance data is required for water reduction review.",
        },
        {
            "discipline": "Landscape",
            "action": "Provide irrigation strategy, landscape water demand assumptions, and any reduced-water design measures.",
            "reason": "Outdoor water use evidence is needed for landscape-related water credits.",
        },
        {
            "discipline": "Architecture",
            "action": "Confirm any water-using equipment or special spaces that may affect baseline and proposed usage.",
            "reason": "Architectural program information can change water demand assumptions.",
        },
        {
            "discipline": "Sustainability",
            "action": "Prepare indoor and outdoor water calculation sheets aligned with the target credit path.",
            "reason": "A consolidated calculation package is required for review and submittal.",
        },
    ],
    "materials": [
        {
            "discipline": "Architecture",
            "action": "Provide finish schedules, material specifications, and product-level sustainability documentation where available.",
            "reason": "Architectural material data supports materials and carbon-related review.",
        },
        {
            "discipline": "Structure",
            "action": "Provide concrete and steel quantity summaries, mix information, and any low-carbon alternatives under consideration.",
            "reason": "Structural materials usually drive embodied carbon impact.",
        },
        {
            "discipline": "Procurement / Cost",
            "action": "Collect EPDs, recycled content declarations, and supplier sustainability documents for priority products.",
            "reason": "Supplier documentation is needed to substantiate material-related claims.",
        },
        {
            "discipline": "Sustainability",
            "action": "Prepare a material compliance tracker showing required evidence by package and responsible consultant.",
            "reason": "Tracking is needed to avoid documentation gaps across disciplines.",
        },
    ],
    "ieq": [
        {
            "discipline": "Mechanical",
            "action": "Provide ventilation calculations, outside air assumptions, filtration approach, and IAQ-related control strategy.",
            "reason": "Ventilation evidence is a key IEQ input.",
        },
        {
            "discipline": "Architecture",
            "action": "Provide daylight, glazing, shading, and spatial planning information relevant to occupied spaces.",
            "reason": "Architectural design decisions strongly affect IEQ performance.",
        },
        {
            "discipline": "Interior Design",
            "action": "Provide low-VOC material specifications and interior finish schedules.",
            "reason": "Interior product selections affect emissions-related IEQ review.",
        },
        {
            "discipline": "Sustainability",
            "action": "Prepare an IEQ evidence matrix covering ventilation, materials, daylight, and thermal comfort inputs.",
            "reason": "A coordinated matrix reduces missing evidence during LEED review.",
        },
    ],
}

ENERGY_READINESS_ITEMS = [
    {
        "readiness_item_id": "envelope_data",
        "readiness_item_name": "Envelope Data",
        "keywords": ["envelope", "u-value", "glazing", "roof", "wall", "shading"],
        "summary": "Building envelope thermal inputs for simulation.",
        "missing_inputs": [
            "Wall/roof/glazing performance values",
            "Window-to-wall ratio or façade assumptions",
            "Shading or solar control information",
        ],
    },
    {
        "readiness_item_id": "hvac_system",
        "readiness_item_name": "HVAC System",
        "keywords": ["hvac", "chiller", "ahu", "fcu", "cop", "cooling", "heating"],
        "summary": "Mechanical system type and performance inputs.",
        "missing_inputs": [
            "System type and zoning logic",
            "Equipment efficiencies",
            "Control sequence and setpoints",
        ],
    },
    {
        "readiness_item_id": "lighting_and_loads",
        "readiness_item_name": "Lighting and Equipment Loads",
        "keywords": ["lighting", "lpd", "equipment", "plug load", "load density"],
        "summary": "Lighting and internal equipment load inputs.",
        "missing_inputs": [
            "Lighting power density",
            "Plug/process load assumptions",
            "Control strategy for lighting",
        ],
    },
    {
        "readiness_item_id": "occupancy_and_schedules",
        "readiness_item_name": "Occupancy and Schedules",
        "keywords": ["occupancy", "schedule", "operation", "hours", "people"],
        "summary": "Occupancy, operating hours, and schedule assumptions.",
        "missing_inputs": [
            "Occupancy density",
            "Operating schedules",
            "Holiday or weekend schedule assumptions",
        ],
    },
    {
        "readiness_item_id": "ventilation",
        "readiness_item_name": "Ventilation / Outside Air",
        "keywords": ["ventilation", "outside air", "fresh air", "ashrae", "air change"],
        "summary": "Ventilation and outside air design basis.",
        "missing_inputs": [
            "Outside air rates",
            "Ventilation calculation basis",
            "Demand control ventilation assumptions",
        ],
    },
    {
        "readiness_item_id": "dhw",
        "readiness_item_name": "Domestic Hot Water",
        "keywords": ["domestic hot water", "dhw", "hot water", "water heater", "boiler"],
        "summary": "Domestic hot water system and related energy inputs.",
        "missing_inputs": [
            "DHW system description",
            "Fuel type and efficiency",
            "Hot water demand assumptions",
        ],
    },
    {
        "readiness_item_id": "baseline_reference",
        "readiness_item_name": "Baseline / Code Reference",
        "keywords": ["baseline", "ashrae", "appendix g", "code", "reference building"],
        "summary": "Baseline model or code reference readiness.",
        "missing_inputs": [
            "Applicable baseline standard",
            "Modeling rules reference",
            "Baseline interpretation narrative",
        ],
    },
    {
        "readiness_item_id": "renewables",
        "readiness_item_name": "Renewables",
        "keywords": ["solar", "pv", "photovoltaic", "renewable", "battery"],
        "summary": "Renewable energy input readiness.",
        "missing_inputs": [
            "PV capacity or area",
            "Renewable generation assumptions",
            "Grid interaction/storage assumptions",
        ],
    },
]

ENERGY_ACTION_LIBRARY = {
    "envelope_data": [
        {
            "discipline": "Architecture",
            "action": "Provide envelope assembly performance data, glazing details, and shading assumptions.",
            "reason": "Envelope thermal properties are required for energy model inputs.",
        },
        {
            "discipline": "Sustainability",
            "action": "Check that envelope inputs are aligned with the intended LEED energy modeling approach.",
            "reason": "Model assumptions must be documented consistently.",
        },
    ],
    "hvac_system": [
        {
            "discipline": "Mechanical",
            "action": "Provide HVAC system narrative, equipment efficiencies, zoning logic, and control sequence.",
            "reason": "Mechanical system inputs are essential for simulation readiness.",
        },
        {
            "discipline": "Sustainability",
            "action": "Prepare an HVAC input sheet for simulation handoff.",
            "reason": "A consistent simulation-ready dataset is needed.",
        },
    ],
    "lighting_and_loads": [
        {
            "discipline": "Electrical",
            "action": "Provide lighting power density, lighting controls, and major internal equipment loads.",
            "reason": "Lighting and plug loads materially affect modeled energy use.",
        },
        {
            "discipline": "Sustainability",
            "action": "Map all load assumptions into a simulation input template.",
            "reason": "Load consistency is needed across all model zones.",
        },
    ],
    "occupancy_and_schedules": [
        {
            "discipline": "Architecture",
            "action": "Confirm space types, occupancy assumptions, and operating patterns.",
            "reason": "Space program and occupancy directly affect schedules.",
        },
        {
            "discipline": "Sustainability",
            "action": "Prepare standardized weekday/weekend/holiday schedules.",
            "reason": "Simulation requires explicit schedule sets.",
        },
    ],
    "ventilation": [
        {
            "discipline": "Mechanical",
            "action": "Provide outside air rates, ventilation method, and any demand-control logic.",
            "reason": "Ventilation assumptions affect fan and conditioning loads.",
        },
        {
            "discipline": "Sustainability",
            "action": "Cross-check ventilation assumptions against the baseline/reference standard.",
            "reason": "Baseline and proposed ventilation assumptions must be consistent.",
        },
    ],
    "dhw": [
        {
            "discipline": "Mechanical",
            "action": "Provide domestic hot water system type, fuel, and efficiency assumptions.",
            "reason": "DHW can be a non-trivial simulation input depending on asset type.",
        },
        {
            "discipline": "Plumbing",
            "action": "Provide hot water demand assumptions or fixture-based demand inputs.",
            "reason": "Demand assumptions are needed for DHW energy modeling.",
        },
    ],
    "baseline_reference": [
        {
            "discipline": "Sustainability",
            "action": "Define the baseline/code framework and document modeling assumptions.",
            "reason": "A clear baseline framework is required before simulation starts.",
        },
        {
            "discipline": "Energy Modeler",
            "action": "Prepare a baseline interpretation note for review.",
            "reason": "The simulation method must be traceable and reviewable.",
        },
    ],
    "renewables": [
        {
            "discipline": "Electrical",
            "action": "Provide PV or renewable system capacity, layout assumptions, and expected generation inputs.",
            "reason": "Renewable systems need explicit modeled assumptions.",
        },
        {
            "discipline": "Sustainability",
            "action": "Clarify whether renewables are part of base scope or improvement scenario.",
            "reason": "Scenario boundaries affect simulation results and reporting.",
        },
    ],
}

CARBON_READINESS_ITEMS = [
    {
        "carbon_item_id": "structural_materials",
        "carbon_item_name": "Structural Materials",
        "keywords": ["concrete", "steel", "cement", "rebar", "tonnage", "mix"],
        "summary": "Structural material quantities and low-carbon opportunities.",
        "missing_inputs": [
            "Concrete volumes or mix classes",
            "Steel tonnage or section schedule",
            "Low-carbon structural alternatives",
        ],
        "decarbonization_actions": [
            "Review lower-cement concrete mixes",
            "Review recycled steel sourcing options",
            "Prioritize high-volume structural packages first",
        ],
    },
    {
        "carbon_item_id": "envelope_materials",
        "carbon_item_name": "Envelope Materials",
        "keywords": ["facade", "cladding", "insulation", "glazing", "aluminum", "envelope"],
        "summary": "Envelope product inputs affecting embodied carbon.",
        "missing_inputs": [
            "Facade system quantities",
            "Insulation and glazing specifications",
            "Alternative low-carbon envelope options",
        ],
        "decarbonization_actions": [
            "Compare façade material alternatives",
            "Review insulation product declarations",
            "Assess glazing specification optimization",
        ],
    },
    {
        "carbon_item_id": "interior_finishes",
        "carbon_item_name": "Interior Finishes",
        "keywords": ["finish", "gypsum", "tile", "paint", "carpet", "ceiling"],
        "summary": "Interior finish package data for embodied carbon review.",
        "missing_inputs": [
            "Finish schedules",
            "Material quantities by area/type",
            "Product substitution options",
        ],
        "decarbonization_actions": [
            "Identify high-volume interior finish packages",
            "Review lower-impact finish alternatives",
            "Collect finish product declarations",
        ],
    },
    {
        "carbon_item_id": "epd_and_transparency",
        "carbon_item_name": "EPD / Product Transparency",
        "keywords": ["epd", "environmental product declaration", "product declaration", "iso 14025"],
        "summary": "Availability of EPD and transparency evidence.",
        "missing_inputs": [
            "EPD documents for priority products",
            "Supplier sustainability declarations",
            "Product-level transparency tracker",
        ],
        "decarbonization_actions": [
            "Prioritize EPD collection for major materials",
            "Create an EPD tracker by package",
            "Engage suppliers on missing declarations",
        ],
    },
    {
        "carbon_item_id": "lca_inputs",
        "carbon_item_name": "LCA Inputs / Quantity Takeoff",
        "keywords": ["lca", "quantity takeoff", "boq", "quantity", "material schedule", "takeoff"],
        "summary": "Quantity and scope inputs required for embodied carbon assessment.",
        "missing_inputs": [
            "BOQ or quantity takeoff",
            "Scope boundaries for assessment",
            "Package-level quantity breakdown",
        ],
        "decarbonization_actions": [
            "Prepare quantity takeoff for major materials",
            "Define assessment scope and boundaries",
            "Map quantities to material categories",
        ],
    },
    {
        "carbon_item_id": "construction_and_waste",
        "carbon_item_name": "Construction / Waste Strategy",
        "keywords": ["waste", "construction waste", "recycling", "diversion", "site waste"],
        "summary": "Construction and waste-related carbon reduction readiness.",
        "missing_inputs": [
            "Waste management assumptions",
            "Recycling/diversion strategy",
            "Construction carbon reduction measures",
        ],
        "decarbonization_actions": [
            "Prepare waste diversion plan",
            "Review recycled content opportunities",
            "Coordinate site waste reporting expectations",
        ],
    },
]

CARBON_ACTION_LIBRARY = {
    "structural_materials": [
        {
            "discipline": "Structure",
            "action": "Provide structural material quantities, concrete mix data, and steel tonnage.",
            "reason": "Structural packages usually dominate embodied carbon.",
        },
        {
            "discipline": "Sustainability",
            "action": "Prepare an embodied carbon hotspot summary for structural packages.",
            "reason": "Carbon reduction should focus first on high-impact structural elements.",
        },
    ],
    "envelope_materials": [
        {
            "discipline": "Architecture",
            "action": "Provide façade, glazing, insulation, and cladding specifications with quantities where possible.",
            "reason": "Envelope materials can significantly affect embodied carbon outcomes.",
        },
        {
            "discipline": "Sustainability",
            "action": "Compare alternative envelope products for lower-carbon options.",
            "reason": "Envelope substitutions may create decarbonization opportunities.",
        },
    ],
    "interior_finishes": [
        {
            "discipline": "Interior Design",
            "action": "Provide finish schedules and identify high-volume finish categories.",
            "reason": "Finish selections can be optimized early for lower-carbon alternatives.",
        },
        {
            "discipline": "Procurement / Cost",
            "action": "Check supplier availability of lower-carbon finish products.",
            "reason": "Alternative product selection depends on supply chain options.",
        },
    ],
    "epd_and_transparency": [
        {
            "discipline": "Procurement / Cost",
            "action": "Collect EPDs and supplier declarations for priority products.",
            "reason": "Carbon documentation depends on supplier transparency data.",
        },
        {
            "discipline": "Sustainability",
            "action": "Maintain an EPD / transparency compliance register.",
            "reason": "A structured tracker is needed for carbon reporting and LEED alignment.",
        },
    ],
    "lca_inputs": [
        {
            "discipline": "Quantity Surveyor",
            "action": "Prepare BOQ-aligned quantity takeoff for major material categories.",
            "reason": "LCA cannot proceed reliably without quantity inputs.",
        },
        {
            "discipline": "Sustainability",
            "action": "Define LCA scope, system boundary, and assessment assumptions.",
            "reason": "Carbon calculations require explicit scope definitions.",
        },
    ],
    "construction_and_waste": [
        {
            "discipline": "Contractor / Construction",
            "action": "Prepare waste diversion and construction-stage reduction strategy.",
            "reason": "Construction practices influence waste-related carbon reduction outcomes.",
        },
        {
            "discipline": "Sustainability",
            "action": "Define reporting approach for waste and construction-stage carbon measures.",
            "reason": "Consistent reporting is needed for review and documentation.",
        },
    ],
}

LEED_SCORING_ITEMS = [
    {
        "category_id": "integrative_process",
        "category_name": "Integrative Process",
        "max_points": 1,
        "keywords": ["integrative", "basis of design", "owner project requirements", "water budget", "early analysis"],
        "review_note": "Early-stage coordination and basis documents for LEED workflow.",
        "required_documents": [
            "Basis of Design",
            "Owner Project Requirements",
            "Early water or energy analysis summary",
        ],
    },
    {
        "category_id": "sustainable_sites",
        "category_name": "Sustainable Sites",
        "max_points": 10,
        "keywords": ["site", "stormwater", "landscape", "heat island", "open space", "light pollution"],
        "review_note": "Site-related documentation supporting sustainable site measures.",
        "required_documents": [
            "Site plan and landscape package",
            "Stormwater strategy",
            "Heat island / exterior lighting information",
        ],
    },
    {
        "category_id": "water_efficiency",
        "category_name": "Water Efficiency",
        "max_points": 11,
        "keywords": ["water", "fixture", "flow rate", "irrigation", "gpm", "gpf"],
        "review_note": "Indoor and outdoor water reduction support package.",
        "required_documents": [
            "Plumbing fixture schedule",
            "Water calculation sheet",
            "Irrigation / landscape water documentation",
        ],
    },
    {
        "category_id": "energy_atmosphere",
        "category_name": "Energy & Atmosphere",
        "max_points": 33,
        "keywords": ["energy", "ashrae", "hvac", "lighting", "model", "commissioning"],
        "review_note": "Energy modeling, system performance, and commissioning-related documentation.",
        "required_documents": [
            "Energy model narrative",
            "HVAC and lighting input summary",
            "Commissioning-related design information",
        ],
    },
    {
        "category_id": "materials_resources",
        "category_name": "Materials & Resources",
        "max_points": 13,
        "keywords": ["material", "epd", "recycled", "waste", "lca", "product declaration"],
        "review_note": "Material transparency, waste, and embodied carbon related documentation.",
        "required_documents": [
            "Material tracker",
            "EPD / supplier declarations",
            "Waste management strategy",
        ],
    },
    {
        "category_id": "indoor_environmental_quality",
        "category_name": "Indoor Environmental Quality",
        "max_points": 16,
        "keywords": ["ventilation", "voc", "daylight", "iaq", "thermal comfort", "co2"],
        "review_note": "Indoor environment and occupant comfort documentation package.",
        "required_documents": [
            "Ventilation basis",
            "Low-emitting materials information",
            "Daylight / comfort support documents",
        ],
    },
    {
        "category_id": "innovation",
        "category_name": "Innovation / Regional Priority",
        "max_points": 6,
        "keywords": ["innovation", "pilot", "regional priority", "education", "exemplary"],
        "review_note": "Innovation path and special LEED narrative readiness.",
        "required_documents": [
            "Innovation strategy memo",
            "Potential pilot or exemplary narrative",
            "Regional priority mapping note",
        ],
    },
]

LEED_ACTION_LIBRARY = {
    "integrative_process": [
        {
            "discipline": "Sustainability",
            "action": "Prepare integrative process memo and consolidate early design analyses.",
            "reason": "LEED review needs traceable early-stage coordination evidence.",
        },
        {
            "discipline": "Project Management",
            "action": "Collect Basis of Design and Owner Project Requirements in one package.",
            "reason": "Core project definition documents are required for structured submittal.",
        },
    ],
    "sustainable_sites": [
        {
            "discipline": "Landscape",
            "action": "Provide planting, irrigation, and open-space related drawings and schedules.",
            "reason": "Landscape data supports multiple site-related credits.",
        },
        {
            "discipline": "Civil / Site",
            "action": "Provide stormwater and site strategy package.",
            "reason": "Site documentation is needed for LEED site category review.",
        },
    ],
    "water_efficiency": [
        {
            "discipline": "Plumbing",
            "action": "Provide fixture flow/flush rates and water calculation support.",
            "reason": "Indoor water reduction cannot be reviewed without plumbing data.",
        },
        {
            "discipline": "Landscape",
            "action": "Provide outdoor water reduction strategy and irrigation assumptions.",
            "reason": "Outdoor water use must be documented separately.",
        },
    ],
    "energy_atmosphere": [
        {
            "discipline": "Mechanical",
            "action": "Provide HVAC narrative, efficiencies, zoning, and commissioning-related design inputs.",
            "reason": "Mechanical design drives major LEED EA documentation.",
        },
        {
            "discipline": "Electrical",
            "action": "Provide lighting power density, controls, and major electrical load assumptions.",
            "reason": "Lighting and electrical loads affect energy review and scoring.",
        },
        {
            "discipline": "Sustainability",
            "action": "Prepare energy model support package and LEED EA documentation matrix.",
            "reason": "Scoring readiness depends on a coordinated documentation package.",
        },
    ],
    "materials_resources": [
        {
            "discipline": "Procurement / Cost",
            "action": "Collect EPDs, recycled content declarations, and supplier documents.",
            "reason": "Material-related points depend on product documentation availability.",
        },
        {
            "discipline": "Sustainability",
            "action": "Prepare a materials compliance tracker and waste strategy summary.",
            "reason": "LEED MR scoring requires organized evidence and package tracking.",
        },
    ],
    "indoor_environmental_quality": [
        {
            "discipline": "Mechanical",
            "action": "Provide ventilation calculations and IAQ strategy documents.",
            "reason": "IEQ review depends on ventilation and indoor air quality evidence.",
        },
        {
            "discipline": "Interior Design",
            "action": "Provide low-emitting material specifications and finish schedules.",
            "reason": "IEQ material credits require finish-level data.",
        },
    ],
    "innovation": [
        {
            "discipline": "Sustainability",
            "action": "Develop innovation credit opportunities and draft supporting narratives.",
            "reason": "Innovation credits usually require explicit strategy and narrative preparation.",
        },
        {
            "discipline": "Project Management",
            "action": "Confirm whether any regional priority or pilot paths should be targeted.",
            "reason": "Targeted strategy is needed before documentation can be assembled.",
        },
    ],
}


app.add_middleware(
    CORSMiddleware,
    allow_origins=[
        "http://localhost:3001",
        "http://127.0.0.1:3001",
    ],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


@app.on_event("startup")
def on_startup():
    Base.metadata.create_all(bind=engine)

    with engine.begin() as connection:
        connection.execute(text("ALTER TABLE documents ADD COLUMN IF NOT EXISTS parse_status VARCHAR DEFAULT 'uploaded'"))
        connection.execute(text("ALTER TABLE documents ADD COLUMN IF NOT EXISTS parse_message VARCHAR"))
        connection.execute(text("ALTER TABLE documents ADD COLUMN IF NOT EXISTS extracted_text TEXT"))
        connection.execute(text("ALTER TABLE documents ADD COLUMN IF NOT EXISTS parsed_at TIMESTAMP"))


def get_db():
    db = SessionLocal()
    try:
        yield db
    finally:
        db.close()


def parse_text_like_file(file_path: Path) -> str:
    suffix = file_path.suffix.lower()

    if suffix in {".txt", ".md"}:
        return file_path.read_text(encoding="utf-8", errors="replace")

    if suffix == ".json":
        raw_text = file_path.read_text(encoding="utf-8", errors="replace")
        try:
            parsed = json.loads(raw_text)
            return json.dumps(parsed, ensure_ascii=False, indent=2)
        except Exception:
            return raw_text

    if suffix == ".csv":
        raw_text = file_path.read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(StringIO(raw_text))
        rows = []
        for row in reader:
            rows.append(" | ".join(row))
        return "\n".join(rows)

    raise ValueError("Unsupported text-like file type")


def parse_pdf_file(file_path: Path) -> str:
    reader = PdfReader(str(file_path))
    pages = []

    for idx, page in enumerate(reader.pages, start=1):
        text_content = page.extract_text() or ""
        pages.append(f"[Page {idx}]\n{text_content.strip()}")

    return "\n\n".join(pages).strip()


def parse_docx_file(file_path: Path) -> str:
    doc = DocxDocument(str(file_path))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

    table_lines = []
    for table in doc.tables:
        for row in table.rows:
            values = [cell.text.strip() for cell in row.cells]
            if any(values):
                table_lines.append(" | ".join(values))

    combined = []
    if paragraphs:
        combined.append("\n".join(paragraphs))
    if table_lines:
        combined.append("\n".join(table_lines))

    return "\n\n".join(combined).strip()


def parse_xlsx_file(file_path: Path) -> str:
    workbook = load_workbook(filename=str(file_path), data_only=True)
    sheets_output = []

    for sheet in workbook.worksheets:
        rows_output = [f"[Sheet: {sheet.title}]"]
        for row in sheet.iter_rows(values_only=True):
            values = ["" if value is None else str(value) for value in row]
            if any(v.strip() for v in values):
                rows_output.append(" | ".join(values))
        sheets_output.append("\n".join(rows_output))

    return "\n\n".join(sheets_output).strip()


def parse_supported_file(document: Document) -> tuple[str, str | None, str | None]:
    file_path = Path(document.file_path)

    if not file_path.exists():
        return "failed", "File not found on disk.", None

    suffix = file_path.suffix.lower()

    if suffix not in SUPPORTED_PARSE_EXTENSIONS:
        return (
            "pending_parser",
            "Parser not available for this file type yet.",
            None,
        )

    try:
        if suffix in TEXT_EXTENSIONS:
            extracted_text = parse_text_like_file(file_path)
        elif suffix == ".pdf":
            extracted_text = parse_pdf_file(file_path)
        elif suffix == ".docx":
            extracted_text = parse_docx_file(file_path)
        elif suffix == ".xlsx":
            extracted_text = parse_xlsx_file(file_path)
        else:
            return "pending_parser", "Parser not available for this file type yet.", None

        extracted_text = extracted_text.strip()

        if not extracted_text:
            return "parsed", "Parsing completed but no readable text was extracted.", ""

        return "parsed", "Parsing completed successfully.", extracted_text

    except Exception as e:
        return "failed", f"Parsing failed: {str(e)}", None


def build_snippet(source_text: str, query: str, radius: int = 120) -> str:
    if not source_text:
        return ""

    lower_source = source_text.lower()
    lower_query = query.lower()
    index = lower_source.find(lower_query)

    if index == -1:
        snippet = source_text[: radius * 2]
        return snippet.strip()

    start = max(0, index - radius)
    end = min(len(source_text), index + len(query) + radius)
    snippet = source_text[start:end].strip()

    if start > 0:
        snippet = "..." + snippet
    if end < len(source_text):
        snippet = snippet + "..."

    return snippet


def collect_topic_evidence(documents: list[Document], keywords: list[str], max_items: int = 5):
    evidences = []
    total_count = 0

    for document in documents:
        extracted_text = document.extracted_text or ""
        lower_text = extracted_text.lower()

        for keyword in keywords:
            keyword_lower = keyword.lower()
            count = lower_text.count(keyword_lower)
            if count > 0:
                total_count += count
                evidences.append(
                    ReviewEvidenceItemResponse(
                        document_id=document.id,
                        original_filename=document.original_filename,
                        keyword=keyword,
                        snippet=build_snippet(extracted_text, keyword),
                    )
                )
                break

        if len(evidences) >= max_items:
            break

    return evidences, total_count


def determine_finding_status(total_count: int) -> str:
    if total_count >= 3:
        return "evidence_found"
    if total_count >= 1:
        return "limited_evidence"
    return "no_evidence"


def get_topic_score(status: str) -> int:
    if status in {"evidence_found", "ready"}:
        return 100
    if status in {"limited_evidence", "partial"}:
        return 50
    return 0


def determine_overall_status(findings: list[ReviewFindingResponse], parsed_document_count: int) -> str:
    if parsed_document_count == 0:
        return "insufficient_documents"

    statuses = [finding.status for finding in findings]

    if all(status == "evidence_found" for status in statuses):
        return "good_initial_coverage"
    if any(status in {"evidence_found", "limited_evidence"} for status in statuses):
        return "partial_coverage"
    return "insufficient_evidence"


def get_priority_for_status(status: str) -> str:
    if status in {"no_evidence", "missing"}:
        return "high"
    if status in {"limited_evidence", "partial"}:
        return "medium"
    return "low"


def build_corrective_actions(topic_id: str, status: str):
    base_actions = CORRECTIVE_ACTION_LIBRARY.get(topic_id, [])
    priority = get_priority_for_status(status)

    corrective_actions = []
    for item in base_actions:
        if status == "no_evidence":
            action_text = f"Immediately provide: {item['action']}"
            reason_text = f"{item['reason']} Current review found no supporting evidence."
        elif status == "limited_evidence":
            action_text = f"Strengthen documentation: {item['action']}"
            reason_text = f"{item['reason']} Current review found only limited evidence."
        else:
            action_text = f"Refine and verify: {item['action']}"
            reason_text = f"{item['reason']} Evidence exists, but should be validated and organized for submission readiness."

        corrective_actions.append(
            CorrectiveActionResponse(
                discipline=item["discipline"],
                priority=priority,
                action=action_text,
                reason=reason_text,
            )
        )

    return corrective_actions


def build_agent1_review(project_id: int, db: Session) -> Agent1ReviewResponse:
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    documents = (
        db.query(Document)
        .filter(Document.project_id == project_id)
        .order_by(Document.id.desc())
        .all()
    )

    parsed_documents = [
        document
        for document in documents
        if document.parse_status == "parsed" and (document.extracted_text or "").strip()
    ]

    findings = []

    for topic in REVIEW_TOPICS:
        evidences, total_count = collect_topic_evidence(parsed_documents, topic["keywords"])
        status = determine_finding_status(total_count)
        corrective_actions = build_corrective_actions(topic["topic_id"], status)
        score = get_topic_score(status)
        max_score = 100
        progress_percent = score

        findings.append(
            ReviewFindingResponse(
                topic_id=topic["topic_id"],
                topic_name=topic["topic_name"],
                status=status,
                score=score,
                max_score=max_score,
                progress_percent=progress_percent,
                evidence_count=total_count,
                searched_keywords=topic["keywords"],
                recommendation=topic["recommendation"],
                evidences=evidences,
                corrective_actions=corrective_actions,
            )
        )

    overall_status = determine_overall_status(findings, len(parsed_documents))
    overall_score = sum(finding.score for finding in findings)
    overall_max_score = sum(finding.max_score for finding in findings) if findings else 0
    overall_progress_percent = int((overall_score / overall_max_score) * 100) if overall_max_score else 0

    return Agent1ReviewResponse(
        project_id=project.id,
        project_name=project.name,
        overall_status=overall_status,
        overall_score=overall_score,
        overall_max_score=overall_max_score,
        overall_progress_percent=overall_progress_percent,
        reviewed_document_count=len(documents),
        parsed_document_count=len(parsed_documents),
        findings=findings,
    )


def determine_energy_status(total_count: int) -> str:
    if total_count >= 3:
        return "ready"
    if total_count >= 1:
        return "partial"
    return "missing"


def build_energy_corrective_actions(readiness_item_id: str, status: str):
    base_actions = ENERGY_ACTION_LIBRARY.get(readiness_item_id, [])
    priority = get_priority_for_status(status)

    actions = []
    for item in base_actions:
        if status == "missing":
            action_text = f"Immediately provide: {item['action']}"
            reason_text = f"{item['reason']} Current energy readiness review found no usable evidence."
        elif status == "partial":
            action_text = f"Complete missing inputs: {item['action']}"
            reason_text = f"{item['reason']} Current energy readiness review found only partial evidence."
        else:
            action_text = f"Validate simulation input: {item['action']}"
            reason_text = f"{item['reason']} Evidence exists, but it should be checked before simulation."
        actions.append(
            CorrectiveActionResponse(
                discipline=item["discipline"],
                priority=priority,
                action=action_text,
                reason=reason_text,
            )
        )
    return actions


def determine_energy_overall_status(findings: list[EnergyFindingResponse], parsed_document_count: int) -> str:
    if parsed_document_count == 0:
        return "insufficient_documents"

    statuses = [finding.status for finding in findings]

    if all(status == "ready" for status in statuses):
        return "ready_for_simulation"
    if any(status in {"ready", "partial"} for status in statuses):
        return "partial_readiness"
    return "not_ready"


def build_agent2_energy_review(project_id: int, db: Session) -> Agent2EnergyReviewResponse:
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    documents = (
        db.query(Document)
        .filter(Document.project_id == project_id)
        .order_by(Document.id.desc())
        .all()
    )

    parsed_documents = [
        document
        for document in documents
        if document.parse_status == "parsed" and (document.extracted_text or "").strip()
    ]

    findings = []

    for item in ENERGY_READINESS_ITEMS:
        evidences, total_count = collect_topic_evidence(parsed_documents, item["keywords"])
        status = determine_energy_status(total_count)
        score = get_topic_score(status)
        max_score = 100
        progress_percent = score
        corrective_actions = build_energy_corrective_actions(item["readiness_item_id"], status)

        findings.append(
            EnergyFindingResponse(
                readiness_item_id=item["readiness_item_id"],
                readiness_item_name=item["readiness_item_name"],
                status=status,
                score=score,
                max_score=max_score,
                progress_percent=progress_percent,
                evidence_count=total_count,
                searched_keywords=item["keywords"],
                summary=item["summary"],
                missing_inputs=item["missing_inputs"],
                evidences=evidences,
                corrective_actions=corrective_actions,
            )
        )

    overall_status = determine_energy_overall_status(findings, len(parsed_documents))
    overall_score = sum(finding.score for finding in findings)
    overall_max_score = sum(finding.max_score for finding in findings) if findings else 0
    overall_progress_percent = int((overall_score / overall_max_score) * 100) if overall_max_score else 0

    return Agent2EnergyReviewResponse(
        project_id=project.id,
        project_name=project.name,
        overall_status=overall_status,
        overall_score=overall_score,
        overall_max_score=overall_max_score,
        overall_progress_percent=overall_progress_percent,
        reviewed_document_count=len(documents),
        parsed_document_count=len(parsed_documents),
        findings=findings,
    )


def determine_carbon_status(total_count: int) -> str:
    if total_count >= 3:
        return "ready"
    if total_count >= 1:
        return "partial"
    return "missing"


def build_carbon_corrective_actions(carbon_item_id: str, status: str):
    base_actions = CARBON_ACTION_LIBRARY.get(carbon_item_id, [])
    priority = get_priority_for_status(status)

    actions = []
    for item in base_actions:
        if status == "missing":
            action_text = f"Immediately provide: {item['action']}"
            reason_text = f"{item['reason']} Current carbon readiness review found no usable evidence."
        elif status == "partial":
            action_text = f"Complete carbon inputs: {item['action']}"
            reason_text = f"{item['reason']} Current carbon readiness review found only partial evidence."
        else:
            action_text = f"Validate decarbonization input: {item['action']}"
            reason_text = f"{item['reason']} Evidence exists, but should be checked before carbon assessment."
        actions.append(
            CorrectiveActionResponse(
                discipline=item["discipline"],
                priority=priority,
                action=action_text,
                reason=reason_text,
            )
        )
    return actions


def determine_carbon_overall_status(findings: list[CarbonFindingResponse], parsed_document_count: int) -> str:
    if parsed_document_count == 0:
        return "insufficient_documents"

    statuses = [finding.status for finding in findings]

    if all(status == "ready" for status in statuses):
        return "ready_for_carbon_assessment"
    if any(status in {"ready", "partial"} for status in statuses):
        return "partial_readiness"
    return "not_ready"


def build_agent3_carbon_review(project_id: int, db: Session) -> Agent3CarbonReviewResponse:
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    documents = (
        db.query(Document)
        .filter(Document.project_id == project_id)
        .order_by(Document.id.desc())
        .all()
    )

    parsed_documents = [
        document
        for document in documents
        if document.parse_status == "parsed" and (document.extracted_text or "").strip()
    ]

    findings = []

    for item in CARBON_READINESS_ITEMS:
        evidences, total_count = collect_topic_evidence(parsed_documents, item["keywords"])
        status = determine_carbon_status(total_count)
        score = get_topic_score(status)
        max_score = 100
        progress_percent = score
        corrective_actions = build_carbon_corrective_actions(item["carbon_item_id"], status)

        findings.append(
            CarbonFindingResponse(
                carbon_item_id=item["carbon_item_id"],
                carbon_item_name=item["carbon_item_name"],
                status=status,
                score=score,
                max_score=max_score,
                progress_percent=progress_percent,
                evidence_count=total_count,
                searched_keywords=item["keywords"],
                summary=item["summary"],
                missing_inputs=item["missing_inputs"],
                decarbonization_actions=item["decarbonization_actions"],
                evidences=evidences,
                corrective_actions=corrective_actions,
            )
        )

    overall_status = determine_carbon_overall_status(findings, len(parsed_documents))
    overall_score = sum(finding.score for finding in findings)
    overall_max_score = sum(finding.max_score for finding in findings) if findings else 0
    overall_progress_percent = int((overall_score / overall_max_score) * 100) if overall_max_score else 0

    return Agent3CarbonReviewResponse(
        project_id=project.id,
        project_name=project.name,
        overall_status=overall_status,
        overall_score=overall_score,
        overall_max_score=overall_max_score,
        overall_progress_percent=overall_progress_percent,
        reviewed_document_count=len(documents),
        parsed_document_count=len(parsed_documents),
        findings=findings,
    )


def determine_leed_doc_status(total_count: int) -> str:
    if total_count >= 3:
        return "ready"
    if total_count >= 1:
        return "partial"
    return "missing"


def get_leed_estimated_points(status: str, max_points: int) -> int:
    if status == "ready":
        return max_points
    if status == "partial":
        partial = int(round(max_points * 0.4))
        return max(1, partial) if max_points > 0 else 0
    return 0


def build_leed_corrective_actions(category_id: str, status: str):
    base_actions = LEED_ACTION_LIBRARY.get(category_id, [])
    priority = get_priority_for_status(status)

    actions = []
    for item in base_actions:
        if status == "missing":
            action_text = f"Immediately prepare: {item['action']}"
            reason_text = f"{item['reason']} Current LEED documentation review found no usable evidence."
        elif status == "partial":
            action_text = f"Complete documentation: {item['action']}"
            reason_text = f"{item['reason']} Current LEED documentation review found only partial evidence."
        else:
            action_text = f"Verify and organize: {item['action']}"
            reason_text = f"{item['reason']} Evidence exists, but should be organized into submittal-ready format."
        actions.append(
            CorrectiveActionResponse(
                discipline=item["discipline"],
                priority=priority,
                action=action_text,
                reason=reason_text,
            )
        )
    return actions


def determine_leed_overall_status(findings: list[LeedScoringFindingResponse], parsed_document_count: int) -> str:
    if parsed_document_count == 0:
        return "insufficient_documents"

    statuses = [finding.status for finding in findings]

    if all(status == "ready" for status in statuses):
        return "good_documentation_readiness"
    if any(status in {"ready", "partial"} for status in statuses):
        return "partial_documentation_readiness"
    return "insufficient_documentation"


def get_estimated_certification_band(points: int) -> str:
    if points >= 80:
        return "Platinum (starter estimate)"
    if points >= 60:
        return "Gold (starter estimate)"
    if points >= 50:
        return "Silver (starter estimate)"
    if points >= 40:
        return "Certified (starter estimate)"
    return "Below Certified (starter estimate)"


def build_agent4_leed_scoring(project_id: int, db: Session) -> Agent4LeedScoringResponse:
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    documents = (
        db.query(Document)
        .filter(Document.project_id == project_id)
        .order_by(Document.id.desc())
        .all()
    )

    parsed_documents = [
        document
        for document in documents
        if document.parse_status == "parsed" and (document.extracted_text or "").strip()
    ]

    findings = []

    for item in LEED_SCORING_ITEMS:
        evidences, total_count = collect_topic_evidence(parsed_documents, item["keywords"])
        status = determine_leed_doc_status(total_count)
        estimated_points = get_leed_estimated_points(status, item["max_points"])
        progress_percent = int((estimated_points / item["max_points"]) * 100) if item["max_points"] else 0
        corrective_actions = build_leed_corrective_actions(item["category_id"], status)

        missing_documents = item["required_documents"] if status == "missing" else (
            item["required_documents"][1:] if status == "partial" and len(item["required_documents"]) > 1 else []
        )

        findings.append(
            LeedScoringFindingResponse(
                category_id=item["category_id"],
                category_name=item["category_name"],
                status=status,
                estimated_points=estimated_points,
                max_points=item["max_points"],
                progress_percent=progress_percent,
                evidence_count=total_count,
                searched_keywords=item["keywords"],
                review_note=item["review_note"],
                required_documents=item["required_documents"],
                missing_documents=missing_documents,
                evidences=evidences,
                corrective_actions=corrective_actions,
            )
        )

    estimated_points = sum(finding.estimated_points for finding in findings)
    total_possible_points = sum(finding.max_points for finding in findings)
    overall_progress_percent = int((estimated_points / total_possible_points) * 100) if total_possible_points else 0
    overall_status = determine_leed_overall_status(findings, len(parsed_documents))
    estimated_certification_band = get_estimated_certification_band(estimated_points)

    return Agent4LeedScoringResponse(
        project_id=project.id,
        project_name=project.name,
        overall_status=overall_status,
        estimated_points=estimated_points,
        total_possible_points=total_possible_points,
        overall_progress_percent=overall_progress_percent,
        estimated_certification_band=estimated_certification_band,
        target_certification="Gold",
        method_note="This is a starter estimate based on document evidence coverage, not an official LEED certification result.",
        reviewed_document_count=len(documents),
        parsed_document_count=len(parsed_documents),
        findings=findings,
    )


@app.get("/")
def read_root():
    return {"message": "Backend is running"}


@app.get("/health")
def health():
    db_status = "unknown"

    try:
        with engine.connect() as connection:
            connection.execute(text("SELECT 1"))
        db_status = "ok"
    except Exception as e:
        db_status = f"error: {str(e)}"

    return {
        "app_status": "ok",
        "db_status": db_status
    }


@app.post("/projects", response_model=ProjectResponse)
def create_project(payload: ProjectCreate, db: Session = Depends(get_db)):
    project = Project(
        name=payload.name,
        description=payload.description
    )
    db.add(project)
    db.commit()
    db.refresh(project)
    return project


@app.get("/projects", response_model=list[ProjectResponse])
def list_projects(db: Session = Depends(get_db)):
    return db.query(Project).order_by(Project.id.asc()).all()


@app.get("/projects/{project_id}", response_model=ProjectResponse)
def get_project(project_id: int, db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")
    return project


@app.post("/projects/{project_id}/documents", response_model=DocumentResponse)
async def upload_document(
    project_id: int,
    file: UploadFile = File(...),
    db: Session = Depends(get_db)
):
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    project_dir = UPLOAD_DIR / f"project_{project_id}"
    project_dir.mkdir(parents=True, exist_ok=True)

    suffix = Path(file.filename).suffix
    stored_filename = f"{uuid4().hex}{suffix}"
    destination = project_dir / stored_filename

    content = await file.read()
    with open(destination, "wb") as buffer:
        buffer.write(content)

    document = Document(
        project_id=project_id,
        original_filename=file.filename,
        stored_filename=stored_filename,
        file_path=str(destination),
        content_type=file.content_type,
        file_size=len(content),
        parse_status="uploaded",
        parse_message="File uploaded successfully. Parsing not started yet.",
    )

    db.add(document)
    db.commit()
    db.refresh(document)

    return document


@app.get("/projects/{project_id}/documents", response_model=list[DocumentResponse])
def list_documents(project_id: int, db: Session = Depends(get_db)):
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    return (
        db.query(Document)
        .filter(Document.project_id == project_id)
        .order_by(Document.id.desc())
        .all()
    )


@app.post("/documents/{document_id}/parse", response_model=DocumentResponse)
def parse_document(document_id: int, db: Session = Depends(get_db)):
    document = db.query(Document).filter(Document.id == document_id).first()
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    document.parse_status = "processing"
    document.parse_message = "Parsing in progress..."
    db.commit()
    db.refresh(document)

    parse_status, parse_message, extracted_text = parse_supported_file(document)

    document.parse_status = parse_status
    document.parse_message = parse_message
    document.extracted_text = extracted_text
    document.parsed_at = datetime.utcnow()

    db.commit()
    db.refresh(document)

    return document


@app.get(
    "/projects/{project_id}/documents/search",
    response_model=list[DocumentSearchResultResponse],
)
def search_project_documents(
    project_id: int,
    query: str = Query(..., min_length=1),
    limit: int = Query(10, ge=1, le=100),
    db: Session = Depends(get_db),
):
    project = db.query(Project).filter(Project.id == project_id).first()
    if not project:
        raise HTTPException(status_code=404, detail="Project not found")

    normalized_query = query.strip().lower()
    if not normalized_query:
        raise HTTPException(status_code=400, detail="Query must not be empty")

    documents = (
        db.query(Document)
        .filter(Document.project_id == project_id)
        .order_by(Document.id.desc())
        .all()
    )

    results = []

    for document in documents:
        filename = document.original_filename or ""
        extracted_text = document.extracted_text or ""

        filename_count = filename.lower().count(normalized_query)
        text_count = extracted_text.lower().count(normalized_query)
        total_count = filename_count + text_count

        if total_count == 0:
            continue

        matched_field = "extracted_text" if text_count >= filename_count and text_count > 0 else "original_filename"
        snippet_source = extracted_text if matched_field == "extracted_text" else filename
        snippet = build_snippet(snippet_source, query)

        results.append(
            DocumentSearchResultResponse(
                document_id=document.id,
                project_id=document.project_id,
                original_filename=document.original_filename,
                parse_status=document.parse_status,
                matched_field=matched_field,
                match_count=total_count,
                snippet=snippet,
            )
        )

    results.sort(key=lambda item: (-item.match_count, item.document_id))
    return results[:limit]


@app.get("/projects/{project_id}/agent1/review", response_model=Agent1ReviewResponse)
def run_agent1_review(project_id: int, db: Session = Depends(get_db)):
    return build_agent1_review(project_id, db)


@app.get("/projects/{project_id}/agent1/export.csv")
def export_agent1_review_csv(project_id: int, db: Session = Depends(get_db)):
    review = build_agent1_review(project_id, db)

    output = StringIO()
    writer = csv.writer(output)

    writer.writerow(["Project ID", review.project_id])
    writer.writerow(["Project Name", review.project_name])
    writer.writerow(["Overall Status", review.overall_status])
    writer.writerow(["Overall Score", f"{review.overall_score}/{review.overall_max_score}"])
    writer.writerow(["Overall Progress", f"{review.overall_progress_percent}%"])
    writer.writerow(["Reviewed Document Count", review.reviewed_document_count])
    writer.writerow(["Parsed Document Count", review.parsed_document_count])
    writer.writerow([])

    writer.writerow([
        "Topic ID",
        "Topic Name",
        "Status",
        "Score",
        "Max Score",
        "Progress Percent",
        "Evidence Count",
        "Searched Keywords",
        "Recommendation",
        "Discipline",
        "Priority",
        "Corrective Action",
        "Reason",
        "Evidence Document",
        "Evidence Keyword",
        "Evidence Snippet",
    ])

    for finding in review.findings:
        max_rows = max(1, len(finding.corrective_actions), len(finding.evidences))

        for i in range(max_rows):
            corrective_action = finding.corrective_actions[i] if i < len(finding.corrective_actions) else None
            evidence = finding.evidences[i] if i < len(finding.evidences) else None

            writer.writerow([
                finding.topic_id if i == 0 else "",
                finding.topic_name if i == 0 else "",
                finding.status if i == 0 else "",
                finding.score if i == 0 else "",
                finding.max_score if i == 0 else "",
                finding.progress_percent if i == 0 else "",
                finding.evidence_count if i == 0 else "",
                ", ".join(finding.searched_keywords) if i == 0 else "",
                finding.recommendation if i == 0 else "",
                corrective_action.discipline if corrective_action else "",
                corrective_action.priority if corrective_action else "",
                corrective_action.action if corrective_action else "",
                corrective_action.reason if corrective_action else "",
                evidence.original_filename if evidence else "",
                evidence.keyword if evidence else "",
                evidence.snippet if evidence else "",
            ])

    output.seek(0)
    filename = f"agent1_review_project_{review.project_id}.csv"

    return StreamingResponse(
        iter([output.getvalue()]),
        media_type="text/csv",
        headers={"Content-Disposition": f'attachment; filename="{filename}"'},
    )


@app.get("/projects/{project_id}/agent2/energy-review", response_model=Agent2EnergyReviewResponse)
def run_agent2_energy_review(project_id: int, db: Session = Depends(get_db)):
    return build_agent2_energy_review(project_id, db)


@app.get("/projects/{project_id}/agent3/carbon-review", response_model=Agent3CarbonReviewResponse)
def run_agent3_carbon_review(project_id: int, db: Session = Depends(get_db)):
    return build_agent3_carbon_review(project_id, db)


@app.get("/projects/{project_id}/agent4/leed-scoring", response_model=Agent4LeedScoringResponse)
def run_agent4_leed_scoring(project_id: int, db: Session = Depends(get_db)):
    return build_agent4_leed_scoring(project_id, db)
